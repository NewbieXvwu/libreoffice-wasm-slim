#!/usr/bin/env python3
"""生成冒烟测试用的 fixture.docx（独立脚本，不依赖 pdf2doc skill）。

内容刻意覆盖本项目的关键渲染路径：
- CJK 正文（验证剪枝后的中文字体回退）
- 加粗 / 下划线 run
- 制表位选项行（A/B/C/D 横排）
- 段落下边框填空线（pBdr）
- 页脚 PAGE / NUMPAGES 域
- Exact 行距

用法: python3 make_fixture.py [输出路径，默认 fixture.docx]
"""
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


def dual_font(run, ascii_font="Times New Roman", cjk="宋体", size_pt=12, bold=False, underline=False):
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn("w:eastAsia"), cjk)
    run.bold = bold
    run.underline = underline


def add_field(paragraph, instr):
    """PAGE / NUMPAGES 域（与 pdf2doc skill 的 helper 同构）。"""
    def mk(kind=None, text=None):
        r = OxmlElement("w:r")
        if kind:
            f = OxmlElement("w:fldChar")
            f.set(qn("w:fldCharType"), kind)
            r.append(f)
        if text is not None:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = f" {instr} "
            r.append(it)
        paragraph._p.append(r)
    mk("begin")
    mk(text=instr)
    mk("end")


def main():
    out = sys.argv[1] if len(sys.argv) > 1 else "fixture.docx"
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = sec.bottom_margin = Mm(25.4)
    sec.left_margin = sec.right_margin = Mm(31.8)

    # 标题：居中加粗
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dual_font(t.add_run("高三英语阶段性检测卷"), size_pt=16, bold=True)

    # CJK 正文段（Exact 行距 16pt）
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(16)
    dual_font(p.add_run("阅读下面的短文，从每题所给的四个选项中选出最佳答案。"
                        "北京市交通管理部门表示，今年将在全市范围内新增五十个智能信号控制路口，"
                        "并根据早晚高峰的车流量动态调整信号灯配时。"))

    # 加粗 + 下划线混排
    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p2.paragraph_format.line_spacing = Pt(16)
    dual_font(p2.add_run("What does the speaker suggest the listeners "))
    dual_font(p2.add_run("pay attention to"), underline=True)
    dual_font(p2.add_run(" during the "))
    dual_font(p2.add_run("morning rush hour"), bold=True)
    dual_font(p2.add_run("?"))

    # 制表位选项行
    opt = doc.add_paragraph()
    opf = opt.paragraph_format
    opf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    opf.line_spacing = Pt(16)
    for mm in (30, 75, 120):
        opf.tab_stops.add_tab_stop(Mm(mm), WD_TAB_ALIGNMENT.LEFT)
    dual_font(opt.add_run("A. To drive carefully\tB. To take the subway\tC. To leave early\tD. To stay home"))

    # 填空线：段落下边框（禁止用 "____" 字符）
    rule = doc.add_paragraph()
    rpf = rule.paragraph_format
    rpf.space_before, rpf.space_after = Pt(9), Pt(11)
    rpf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    rpf.line_spacing = Pt(12)
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 页脚：第 X 页（共 Y 页）域
    footer = sec.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dual_font(fp.add_run("第 "), size_pt=9)
    add_field(fp, "PAGE")
    dual_font(fp.add_run(" 页（共 "), size_pt=9)
    add_field(fp, "NUMPAGES")
    dual_font(fp.add_run(" 页）"), size_pt=9)

    doc.save(out)
    print(f"fixture written: {out}")


if __name__ == "__main__":
    main()
